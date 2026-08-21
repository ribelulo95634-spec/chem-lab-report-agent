import io
import json
import unittest

from docx import Document

from web_app_utils import (
    CATALOG,
    analyze_guide_text,
    build_project_package,
    build_formal_report,
    build_prelab_report,
    build_standard_report_template,
    extract_structured_data_file,
    extract_uploaded_document,
    format_data_rows_markdown,
    load_entry,
    inspect_report_readiness,
    markdown_to_docx_bytes,
    load_project_package,
)


class WebAppUtilsTests(unittest.TestCase):
    def test_catalog_is_examples_not_the_general_input_limit(self):
        for key, entry in CATALOG.items():
            self.assertTrue(entry.spec_path.exists(), key)
            _, spec, report, audit = load_entry(key)
            self.assertIn("experiment_name", spec)
            self.assertIn("预习报告", report)
            self.assertIn("status", audit)

    def test_docx_upload_is_extracted_in_memory(self):
        document = Document()
        document.add_heading("酸碱滴定", 1)
        document.add_heading("实验目的", 2)
        document.add_paragraph("掌握滴定操作")
        document.add_heading("实验步骤", 2)
        document.add_paragraph("记录滴定管初读数和终读数")
        payload = io.BytesIO()
        document.save(payload)

        source = extract_uploaded_document("guide.docx", payload.getvalue())
        self.assertEqual(source["method"], "python_docx")
        self.assertIn("掌握滴定操作", source["text"])
        self.assertEqual(len(source["sha256"]), 64)

    def test_general_analysis_and_prelab_keep_missing_visible(self):
        source = extract_uploaded_document(
            "guide.md",
            "# 酸碱滴定\n\n## 实验目的\n掌握滴定\n\n## 实验步骤\n记录读数".encode("utf-8"),
        )
        analysis = analyze_guide_text(source["text"], "酸碱滴定")
        report = build_prelab_report(analysis, source)
        self.assertEqual(analysis["experiment_name"], "酸碱滴定")
        self.assertIn("掌握滴定", report)
        self.assertIn("【指导书中未识别到", report)
        self.assertNotIn("姓名", report)
        self.assertNotIn("学号", report)

    def test_formal_report_never_invents_missing_results(self):
        source = extract_uploaded_document("guide.txt", "实验目的\n观察反应\n实验步骤\n完成操作".encode("utf-8"))
        analysis = analyze_guide_text(source["text"], "示例实验")
        report = build_formal_report(
            analysis,
            source,
            actual_procedure="",
            observations="",
            raw_data="",
            software_results="",
            error_analysis="",
            conclusion="",
        )
        self.assertIn("【缺失：请填写本次真实观察", report)
        self.assertIn("【缺失：请填写结构化原始数据", report)
        self.assertIn("NEEDS_HUMAN_REVIEW", report)

    def test_csv_data_import_accepts_common_headers(self):
        imported = extract_structured_data_file(
            "data.csv",
            "样品,Unit,Value,说明\n样品A,mL,24.30,初始读数\n".encode("utf-8"),
        )
        self.assertEqual(imported["method"], "csv")
        self.assertEqual(
            imported["rows"],
            [{"数据项": "样品A", "单位": "mL", "原始值": "24.30", "备注": "初始读数"}],
        )

    def test_xlsx_data_import_and_markdown_output(self):
        from openpyxl import Workbook

        workbook = Workbook()
        sheet = workbook.active
        sheet.append(["数据项", "单位", "原始值", "备注"])
        sheet.append(["吸光度", "AU", 0.456, "样品|1"])
        output = io.BytesIO()
        workbook.save(output)
        workbook.close()

        imported = extract_structured_data_file("results.xlsx", output.getvalue())
        markdown = format_data_rows_markdown(imported["rows"])
        self.assertEqual(imported["method"], "xlsx")
        self.assertIn("| 吸光度 | AU | 0.456 | 样品\\|1 |", markdown)

        document = Document(io.BytesIO(markdown_to_docx_bytes("# 数据报告\n\n" + markdown)))
        self.assertEqual(document.tables[0].rows[1].cells[3].text, "样品|1")

    def test_report_readiness_identifies_missing_user_evidence(self):
        readiness = inspect_report_readiness(
            actual_procedure="完成滴定",
            observations="溶液变为浅红色",
            raw_data="24.30 mL",
            software_results="线性拟合 R2=0.998",
            error_analysis="终点判断存在读数误差",
            conclusion="结果满足要求",
        )
        self.assertTrue(readiness["is_ready"])
        self.assertEqual(readiness["completed_count"], 6)

        missing = inspect_report_readiness(
            actual_procedure="",
            observations="",
            raw_data="1.0 g",
            software_results="",
            error_analysis="",
            conclusion="",
        )
        self.assertFalse(missing["is_ready"])
        self.assertIn("真实实验现象", missing["missing"])

        structured = inspect_report_readiness(
            actual_procedure="完成滴定",
            observations="溶液变为浅红色",
            raw_data="",
            software_results="线性拟合 R2=0.998",
            error_analysis="终点判断存在读数误差",
            conclusion="结果满足要求",
            data_rows=[{"数据项": "滴定体积", "单位": "mL", "原始值": "24.30", "备注": ""}],
        )
        self.assertTrue(structured["is_ready"])

    def test_project_package_round_trip_restores_progress(self):
        source = extract_uploaded_document(
            "guide.md",
            "# 酸碱滴定\n\n## 实验目的\n掌握滴定".encode("utf-8"),
        )
        analysis = analyze_guide_text(source["text"], "酸碱滴定")
        payload = build_project_package(
            source,
            analysis,
            {"raw_data": "24.30 mL", "observations": "浅红色保持 30 秒"},
            [{"数据项": "滴定体积", "单位": "mL", "原始值": "24.30", "备注": "第1次"}],
        )
        restored = load_project_package(payload)
        self.assertEqual(restored["source"]["sha256"], source["sha256"])
        self.assertEqual(restored["analysis"]["experiment_name"], "酸碱滴定")
        self.assertEqual(restored["formal_inputs"]["raw_data"], "24.30 mL")
        self.assertEqual(restored["formal_inputs"]["observations"], "浅红色保持 30 秒")
        self.assertEqual(restored["data_rows"][0]["原始值"], "24.30")

    def test_version_one_project_migrates_without_structured_rows(self):
        source = extract_uploaded_document("guide.txt", "实验目的：测试".encode("utf-8"))
        analysis = analyze_guide_text(source["text"], "迁移测试")
        package = json.loads(build_project_package(source, analysis).decode("utf-8"))
        package["schema_version"] = "chemreport.project.v0.1"
        package.pop("data_rows")

        restored = load_project_package(json.dumps(package, ensure_ascii=False).encode("utf-8"))
        self.assertEqual(restored["schema_version"], "chemreport.project.v0.2")
        self.assertEqual(restored["data_rows"], [])

    def test_project_package_rejects_unknown_version(self):
        with self.assertRaisesRegex(ValueError, "版本不受支持"):
            load_project_package(b'{"schema_version":"unknown"}')

    def test_legacy_doc_has_actionable_error(self):
        with self.assertRaisesRegex(ValueError, "另存为 .docx"):
            extract_uploaded_document("guide.doc", b"legacy")

    def test_search_known_experiment_builds_reference_template(self):
        report, reference = build_standard_report_template("酸碱标准溶液标定")
        self.assertIsNotNone(reference)
        self.assertEqual(reference["name"], "酸碱标准溶液标定与未知酸测定")
        self.assertIn("标准实验报告模板", report)
        self.assertIn("原始数据", report)
        self.assertIn("本校实验参数", report)

    def test_search_unknown_experiment_stays_generic(self):
        report, reference = build_standard_report_template("自定义催化实验")
        self.assertIsNone(reference)
        self.assertIn("未在项目实验目录中找到匹配项", report)
        self.assertIn("不根据实验名称猜测", report)

    def test_docx_download_is_created_in_memory(self):
        payload = markdown_to_docx_bytes("# 示例实验 - 预习报告\n\n## 目的\n\n核对流程")
        self.assertGreater(len(payload), 10000)
        self.assertEqual(payload[:2], b"PK")


if __name__ == "__main__":
    unittest.main()
