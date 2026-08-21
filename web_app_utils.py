"""Pure helpers for the public Streamlit chemistry report assistant."""

from __future__ import annotations

import csv
import hashlib
import io
import json
import re
import tempfile
from difflib import get_close_matches
from dataclasses import dataclass
from pathlib import Path


ROOT = Path(__file__).resolve().parent
MAX_SOURCE_CHARS = 120_000
MAX_PROJECT_PACKAGE_BYTES = 5 * 1024 * 1024
PROJECT_SCHEMA_VERSION = "chemreport.project.v0.2"
SUPPORTED_PROJECT_SCHEMA_VERSIONS = {"chemreport.project.v0.1", PROJECT_SCHEMA_VERSION}
DATA_COLUMNS = ("数据项", "单位", "原始值", "备注")
DATA_COLUMN_ALIASES = {
    "数据项": {"数据项", "项目", "名称", "样品", "item", "name", "sample"},
    "单位": {"单位", "unit"},
    "原始值": {"原始值", "数值", "值", "value", "rawvalue", "raw_value"},
    "备注": {"备注", "说明", "note", "notes", "comment"},
}
FORMAL_INPUT_FIELDS = (
    "actual_procedure",
    "observations",
    "raw_data",
    "software_results",
    "error_analysis",
    "conclusion",
)

COMMON_EXPERIMENTS = [
    "容量器皿使用与标准溶液配制",
    "醋酸电离常数与电离度测定",
    "化学反应速率与温度影响",
    "硝酸钾的制备与纯化",
    "硫酸亚铁铵的制备",
    "电离平衡、沉淀平衡与沉淀转化",
    "酸碱标准溶液标定与未知酸测定",
    "EDTA 配位滴定测定水硬度",
    "高锰酸钾法测定铁或过氧化氢",
    "重结晶与熔点测定",
    "常压蒸馏与沸点测定",
    "乙酰水杨酸的制备与纯化",
    "乙酸乙酯皂化反应速率常数的测定",
    "最大泡压法测定表面张力",
    "二组分固液相图测绘",
    "紫外可见分光光度法测定铁",
    "离子选择电极测定氟离子",
    "气相色谱分析混合醇",
]


@dataclass(frozen=True)
class ExperimentEntry:
    key: str
    label: str
    spec_path: Path
    report_path: Path
    audit_path: Path | None
    short_description: str


CATALOG = {
    "liquid_liquid_equilibrium": ExperimentEntry(
        key="liquid_liquid_equilibrium",
        label="乙醇-环己烷-水三元液液平衡",
        spec_path=ROOT / "experiments" / "liquid_liquid_equilibrium" / "config" / "experiment_spec.v0.2.json",
        report_path=ROOT / "outputs" / "reports" / "ethanol_cyclohexane_water_lle" / "实验预习报告.md",
        audit_path=ROOT / "outputs" / "reports" / "ethanol_cyclohexane_water_lle" / "报告审查.json",
        short_description="真实课程讲义接入样例，包含浊点、联结线和 Origin 三元相图计划。",
    ),
    "uv_vis_iron": ExperimentEntry(
        key="uv_vis_iron",
        label="邻菲啰啉分光光度法测定水样中铁",
        spec_path=ROOT / "experiments" / "uv_vis_iron" / "config" / "experiment_spec.water_project_guide.v0.2.json",
        report_path=ROOT / "outputs" / "report_first_demo" / "uv_vis_iron" / "预习报告.md",
        audit_path=ROOT / "outputs" / "report_first_demo" / "uv_vis_iron" / "报告审查.json",
        short_description="报告优先版演示，包含来源追溯、线性拟合计划和 Origin 协作。",
    ),
}


SECTION_RULES = {
    "objectives": ("实验目的", "目的与要求", "实验要求"),
    "principle": ("实验原理", "基本原理", "原理"),
    "materials": ("仪器与试剂", "仪器试剂", "实验仪器", "主要仪器", "药品", "试剂"),
    "safety": ("安全", "注意事项", "废液", "废物处理"),
    "procedure": ("实验步骤", "实验方法", "实验内容", "操作步骤", "操作方法"),
    "data": ("数据处理", "数据记录", "结果处理", "结果与讨论", "作图"),
    "questions": ("思考题", "讨论题", "问题"),
}


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def load_entry(key: str) -> tuple[ExperimentEntry, dict, str, dict]:
    entry = CATALOG[key]
    spec = load_json(entry.spec_path)
    report = entry.report_path.read_text(encoding="utf-8")
    audit = load_json(entry.audit_path) if entry.audit_path and entry.audit_path.exists() else {"status": "NEEDS_HUMAN_REVIEW", "issues": []}
    return entry, spec, report, audit


def safe_text(value: str, fallback: str = "未填写") -> str:
    cleaned = re.sub(r"[\r\n|]+", " ", str(value).strip())
    return cleaned or fallback


def normalize_data_rows(rows) -> list[dict[str, str]]:
    """Normalize tabular experiment records to the four public columns."""

    if rows is None:
        return []
    if hasattr(rows, "to_dict"):
        rows = rows.to_dict("records")
    if not isinstance(rows, list):
        raise ValueError("结构化数据必须是表格行列表。")
    normalized: list[dict[str, str]] = []
    for raw_row in rows[:500]:
        if not isinstance(raw_row, dict):
            continue
        row = {
            column: str(raw_row.get(column, "")).strip()[:2000]
            for column in DATA_COLUMNS
        }
        if any(row.values()):
            normalized.append(row)
    return normalized


def _canonical_data_column(column: str) -> str | None:
    normalized = re.sub(r"[\s\-]+", "", str(column).strip()).lower()
    for canonical, aliases in DATA_COLUMN_ALIASES.items():
        if normalized in {re.sub(r"[\s\-]+", "", alias).lower() for alias in aliases}:
            return canonical
    return None


def _rows_from_grid(values: list[list[object]]) -> list[dict[str, str]]:
    if not values:
        raise ValueError("数据文件中没有可读取的内容。")
    headers = [str(value).strip() if value is not None else "" for value in values[0]]
    mapped = [_canonical_data_column(header) for header in headers]
    if not any(mapped):
        if len(headers) > len(DATA_COLUMNS):
            raise ValueError("未识别表头，请使用：数据项、单位、原始值、备注。")
        mapped = list(DATA_COLUMNS[: len(headers)])
    rows: list[dict[str, str]] = []
    for values_row in values[1:501]:
        row = {column: "" for column in DATA_COLUMNS}
        for index, value in enumerate(values_row):
            if index >= len(mapped) or mapped[index] is None:
                continue
            row[mapped[index]] = "" if value is None else str(value).strip()
        rows.append(row)
    normalized = normalize_data_rows(rows)
    if not normalized:
        raise ValueError("数据文件中没有有效数据行。")
    return normalized


def extract_structured_data_file(file_name: str, data: bytes) -> dict:
    """Extract a CSV or XLSX experiment data table in memory."""

    if not data:
        raise ValueError("数据文件为空。")
    if len(data) > 5 * 1024 * 1024:
        raise ValueError("数据文件超过 5 MB，请压缩或拆分后重试。")
    suffix = Path(file_name).suffix.lower()
    if suffix == ".csv":
        text = _decode_text(data)
        values = list(csv.reader(io.StringIO(text)))
        rows, method = _rows_from_grid(values), "csv"
    elif suffix == ".xlsx":
        from openpyxl import load_workbook

        workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        sheet = workbook.active
        values = [list(row) for row in sheet.iter_rows(values_only=True)]
        workbook.close()
        rows, method = _rows_from_grid(values), "xlsx"
    else:
        raise ValueError("暂支持 CSV 和 XLSX 数据文件。")
    return {"file_name": Path(file_name).name, "method": method, "rows": rows}


def build_data_template_csv() -> bytes:
    output = io.StringIO()
    writer = csv.writer(output, lineterminator="\n")
    writer.writerow(DATA_COLUMNS)
    writer.writerow(["样品1", "mL", "", "保留原始读数"])
    return ("\ufeff" + output.getvalue()).encode("utf-8")


def format_data_rows_markdown(rows) -> str:
    normalized = normalize_data_rows(rows)
    if not normalized:
        return ""

    def cell(value: str) -> str:
        return value.replace("|", "\\|").replace("\r", " ").replace("\n", " ")

    lines = [
        "| 数据项 | 单位 | 原始值 | 备注 |",
        "| --- | --- | --- | --- |",
    ]
    lines.extend(
        "| " + " | ".join(cell(row[column]) for column in DATA_COLUMNS) + " |"
        for row in normalized
    )
    return "\n".join(lines)


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "gb18030", "utf-16"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("文本编码无法识别，请将文件另存为 UTF-8、DOCX 或 PDF。")


def _extract_docx(data: bytes) -> str:
    from docx import Document

    document = Document(io.BytesIO(data))
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name if paragraph.style else ""
        if style.startswith("Heading"):
            match = re.search(r"(\d+)$", style)
            level = max(1, min(6, int(match.group(1)))) if match else 2
            blocks.append(f"{'#' * level} {text}")
        else:
            blocks.append(text)
    for table_index, table in enumerate(document.tables, start=1):
        rows = [[" ".join(cell.text.split()).replace("|", "\\|") for cell in row.cells] for row in table.rows]
        if not rows:
            continue
        blocks.extend([
            f"## 表 {table_index}",
            "| " + " | ".join(rows[0]) + " |",
            "| " + " | ".join("---" for _ in rows[0]) + " |",
        ])
        blocks.extend("| " + " | ".join(row) + " |" for row in rows[1:])
    return "\n\n".join(blocks)


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(f"## 第 {index} 页\n\n{text}")
    if not pages:
        raise ValueError("PDF 没有可提取的文字。扫描版文件需要先进行 OCR。")
    return "\n\n".join(pages)


def extract_uploaded_document(file_name: str, data: bytes) -> dict:
    """Extract an uploaded guide in memory and return traceability metadata."""

    if not data:
        raise ValueError("上传文件为空。")
    if len(data) > 10 * 1024 * 1024:
        raise ValueError("文件超过 10 MB，请压缩或拆分后重试。")
    suffix = Path(file_name).suffix.lower()
    if suffix in {".md", ".txt"}:
        text, method = _decode_text(data), "direct_text"
    elif suffix == ".docx":
        text, method = _extract_docx(data), "python_docx"
    elif suffix == ".pdf":
        text, method = _extract_pdf(data), "pypdf"
    elif suffix == ".doc":
        raise ValueError("旧版 Word .doc 无法在网页端可靠解析，请在 Word/WPS 中另存为 .docx 或 PDF。")
    else:
        raise ValueError("暂支持 DOCX、PDF、Markdown 和 TXT 文件。")
    text = text.strip()
    if not text:
        raise ValueError("文件中没有提取到可用文字。")
    if len(text) > MAX_SOURCE_CHARS:
        text = text[:MAX_SOURCE_CHARS] + "\n\n【后续内容因网页处理上限未载入】"
    return {
        "file_name": Path(file_name).name,
        "sha256": hashlib.sha256(data).hexdigest(),
        "method": method,
        "text": text,
    }


def _canonical_section(heading: str) -> str | None:
    normalized = re.sub(r"[\s:：、.．0-9一二三四五六七八九十()（）]+", "", heading)
    for key, aliases in SECTION_RULES.items():
        if any(alias.replace(" ", "") in normalized for alias in aliases):
            return key
    return None


def analyze_guide_text(source_text: str, fallback_name: str = "未命名实验") -> dict:
    """Create a conservative structured draft without inventing missing content."""

    lines = [line.strip() for line in source_text.replace("\r", "").split("\n")]
    nonempty = [line for line in lines if line]
    title = fallback_name
    for line in nonempty[:20]:
        candidate = re.sub(r"^#{1,6}\s*", "", line).strip()
        if 2 <= len(candidate) <= 60 and not candidate.startswith("|"):
            title = candidate
            break

    sections: dict[str, list[str]] = {key: [] for key in SECTION_RULES}
    current: str | None = None
    for line in lines:
        if not line:
            if current and sections[current] and sections[current][-1] != "":
                sections[current].append("")
            continue
        heading = re.sub(r"^#{1,6}\s*", "", line).strip()
        looks_like_heading = bool(re.match(r"^#{1,6}\s+", line)) or len(heading) <= 40
        detected = _canonical_section(heading) if looks_like_heading else None
        if detected:
            current = detected
            continue
        if current:
            sections[current].append(line)

    cleaned = {key: "\n".join(value).strip() for key, value in sections.items()}
    missing = [key for key, value in cleaned.items() if not value]
    return {
        "experiment_name": safe_text(title, fallback_name),
        "sections": cleaned,
        "missing_sections": missing,
        "source_excerpt": "\n".join(nonempty[:80])[:8000],
    }


def _section(value: str, fallback: str = "【指导书中未识别到，请人工补充】") -> str:
    return value.strip() or fallback


def inspect_report_readiness(
    *,
    actual_procedure: str,
    observations: str,
    raw_data: str,
    software_results: str,
    error_analysis: str,
    conclusion: str,
    data_rows=None,
) -> dict:
    """Summarize missing user-supplied evidence before report export."""

    structured_rows = normalize_data_rows(data_rows)
    required = {
        "实际操作与计划偏差": actual_procedure,
        "真实实验现象": observations,
        "原始数据": raw_data or ("structured_data" if structured_rows else ""),
        "计算结果或专业软件产物": software_results,
        "误差分析": error_analysis,
        "实验结论": conclusion,
    }
    missing = [label for label, value in required.items() if not value.strip()]
    return {
        "is_ready": not missing,
        "completed_count": len(required) - len(missing),
        "total_count": len(required),
        "missing": missing,
    }


def build_project_package(
    source: dict,
    analysis: dict,
    formal_inputs: dict | None = None,
    data_rows=None,
) -> bytes:
    """Build a portable project file so a user can continue later."""

    inputs = formal_inputs or {}
    package = {
        "schema_version": PROJECT_SCHEMA_VERSION,
        "source": {
            "file_name": str(source.get("file_name", "未命名指导书")),
            "sha256": str(source.get("sha256", "")),
            "method": str(source.get("method", "unknown")),
            "text": str(source.get("text", "")),
        },
        "analysis": {
            "experiment_name": str(analysis.get("experiment_name", "未命名实验")),
            "sections": {
                key: str(analysis.get("sections", {}).get(key, ""))
                for key in SECTION_RULES
            },
            "source_excerpt": str(analysis.get("source_excerpt", "")),
        },
        "formal_inputs": {
            key: str(inputs.get(key, ""))
            for key in FORMAL_INPUT_FIELDS
        },
        "data_rows": normalize_data_rows(data_rows),
    }
    return json.dumps(package, ensure_ascii=False, indent=2).encode("utf-8")


def load_project_package(data: bytes) -> dict:
    """Validate and normalize a portable ChemReport project file."""

    if not data:
        raise ValueError("项目文件为空。")
    if len(data) > MAX_PROJECT_PACKAGE_BYTES:
        raise ValueError("项目文件超过 5 MB，请确认文件是否来自 ChemReport。")
    try:
        package = json.loads(data.decode("utf-8-sig"))
    except (UnicodeDecodeError, json.JSONDecodeError) as exc:
        raise ValueError("项目文件不是有效的 UTF-8 JSON。") from exc
    if not isinstance(package, dict) or package.get("schema_version") not in SUPPORTED_PROJECT_SCHEMA_VERSIONS:
        raise ValueError("项目文件版本不受支持，请使用 ChemReport 导出的项目文件。")

    source = package.get("source")
    analysis = package.get("analysis")
    if not isinstance(source, dict) or not isinstance(analysis, dict):
        raise ValueError("项目文件缺少指导书或解析结果。")
    source_text = source.get("text")
    source_hash = source.get("sha256")
    if not isinstance(source_text, str) or not source_text.strip():
        raise ValueError("项目文件中没有可恢复的指导书文字。")
    if not isinstance(source_hash, str) or not re.fullmatch(r"[0-9a-fA-F]{64}", source_hash):
        raise ValueError("项目文件中的来源哈希无效。")

    sections = analysis.get("sections")
    if not isinstance(sections, dict):
        raise ValueError("项目文件中的章节结构无效。")
    normalized_sections = {
        key: str(sections.get(key, ""))
        for key in SECTION_RULES
    }
    normalized_analysis = {
        "experiment_name": str(analysis.get("experiment_name", "未命名实验")),
        "sections": normalized_sections,
        "missing_sections": [key for key, value in normalized_sections.items() if not value.strip()],
        "source_excerpt": str(analysis.get("source_excerpt", ""))[:8000],
    }
    inputs = package.get("formal_inputs") if isinstance(package.get("formal_inputs"), dict) else {}
    return {
        "schema_version": PROJECT_SCHEMA_VERSION,
        "source": {
            "file_name": Path(str(source.get("file_name", "未命名指导书"))).name,
            "sha256": source_hash.lower(),
            "method": str(source.get("method", "restored_project")),
            "text": source_text[:MAX_SOURCE_CHARS],
        },
        "analysis": normalized_analysis,
        "formal_inputs": {
            key: str(inputs.get(key, ""))
            for key in FORMAL_INPUT_FIELDS
        },
        "data_rows": normalize_data_rows(package.get("data_rows", [])),
    }


def build_prelab_report(analysis: dict, source: dict) -> str:
    sections = analysis["sections"]
    missing = "、".join(analysis["missing_sections"]) or "无"
    return f"""# {safe_text(analysis['experiment_name'])} - 实验预习报告草稿

> 状态：NEEDS_HUMAN_REVIEW。本报告只整理上传文件中的可访问文字，不补写未出现的配方、公式、剂量、安全参数或实验结果。

## 来源与解析记录

| 项目 | 内容 |
| --- | --- |
| 上传文件 | {safe_text(source['file_name'])} |
| 文件哈希 | `{source['sha256']}` |
| 提取方式 | {source['method']} |
| 未识别章节 | {safe_text(missing, '无')} |

## 实验目的

{_section(sections['objectives'])}

## 实验原理

{_section(sections['principle'])}

## 仪器与试剂

{_section(sections['materials'])}

## 安全、废物与停止条件

{_section(sections['safety'])}

## 计划实验步骤

{_section(sections['procedure'])}

## 原始数据记录表

| 数据项 | 单位 | 原始值 | 记录说明 |
| --- | --- | --- | --- |
| 【按指导书补充】 |  |  | 不得填写预期结果 |

## 数据处理与作图计划

{_section(sections['data'])}

> 复杂拟合、光谱、色谱或模拟结果应由 Origin 或相应专业软件处理，再将参数表和图表导入正式报告。

## 思考题

{_section(sections['questions'], '【指导书中未识别到思考题】')}

## 实验前核对

- [ ] 关键剂量、温度、时间和仪器参数已在原文件中逐项核对
- [ ] 安全条件和废物去向已确认
- [ ] 数据表、单位、公式和作图要求已准备
- [ ] 预期现象没有被当成真实实验结果
"""


def build_formal_report(
    analysis: dict,
    source: dict,
    *,
    actual_procedure: str,
    observations: str,
    raw_data: str,
    software_results: str,
    error_analysis: str,
    conclusion: str,
    data_rows=None,
) -> str:
    sections = analysis["sections"]
    readiness = inspect_report_readiness(
        actual_procedure=actual_procedure,
        observations=observations,
        raw_data=raw_data,
        software_results=software_results,
        error_analysis=error_analysis,
        conclusion=conclusion,
        data_rows=data_rows,
    )
    missing = readiness["missing"]
    status = "NEEDS_HUMAN_REVIEW"
    missing_text = "、".join(missing) or "无，但仍需人工审核"
    structured_data = format_data_rows_markdown(data_rows)
    raw_data_section = structured_data or "【缺失：请填写结构化原始数据，或在补充说明中粘贴原始记录】"
    raw_data_notes = raw_data.strip() or "【无补充说明】"
    return f"""# {safe_text(analysis['experiment_name'])} - 正式实验报告草稿

> 状态：{status}。真实实验内容只来自用户本次填写或外部软件导出，不根据指导书推测结果。

## 追溯信息

| 项目 | 内容 |
| --- | --- |
| 指导书 | {safe_text(source['file_name'])} |
| 文件哈希 | `{source['sha256']}` |
| 尚缺内容 | {safe_text(missing_text)} |

## 实验目的

{_section(sections['objectives'])}

## 实验原理

{_section(sections['principle'])}

## 仪器与试剂

{_section(sections['materials'])}

## 实际操作与计划偏差

{_section(actual_procedure, '【缺失：请填写本次真实操作以及与指导书不同之处】')}

## 实验现象

{_section(observations, '【缺失：请填写本次真实观察，不能使用预期现象替代】')}

## 原始数据

{raw_data_section}

### 原始数据补充说明

```text
{raw_data_notes}
```

## 数据处理与专业软件产物

### 指导书要求

{_section(sections['data'])}

### 本次分析结果

{_section(software_results, '【缺失：请填写计算结果，或粘贴 Origin 等软件的参数表和图表说明】')}

## 误差分析

{_section(error_analysis, '【缺失：应结合本次操作、数据和仪器证据分析，不写通用套话】')}

## 讨论

{_section(sections['questions'], '【待结合思考题和本次结果填写】')}

## 结论

{_section(conclusion, '【缺失：结论必须由本次真实数据支持】')}

## 最终审核

- [ ] 原始数据与正文一致
- [ ] 单位、有效数字、公式输入和图表坐标已核对
- [ ] 软件参数表、图表和正文来自同一批数据
- [ ] 所有缺失项已处理
- [ ] 使用者已完成最终人工确认
"""


def search_experiment_reference(query: str) -> dict | None:
    """Search the project reference catalogue and return a conservative summary."""

    cleaned_query = re.sub(r"\s+", "", query.strip()).lower()
    if not cleaned_query:
        return None
    normalized = {re.sub(r"\s+", "", name).lower(): name for name in COMMON_EXPERIMENTS}
    matched_name = next(
        (name for key, name in normalized.items() if cleaned_query in key or key in cleaned_query),
        None,
    )
    if matched_name is None:
        close = get_close_matches(cleaned_query, list(normalized), n=1, cutoff=0.55)
        matched_name = normalized[close[0]] if close else None
    if matched_name is None:
        return None

    reference_path = ROOT / "docs" / "大学化学实验模块与标准流程.md"
    text = reference_path.read_text(encoding="utf-8")
    pattern = re.compile(
        rf"^###\s+[A-F]\d+\.\s+{re.escape(matched_name)}\s*$([\s\S]*?)(?=^###\s+[A-F]\d+\.|^##\s+|\Z)",
        re.MULTILINE,
    )
    match = pattern.search(text)
    if not match:
        return {"name": matched_name, "objective": "", "procedure": "", "data": "", "matched": True}
    block = match.group(1)

    def bold_field(label: str, next_label: str | None = None) -> str:
        if next_label:
            field_pattern = rf"\*\*{label}\*\*:\s*([\s\S]*?)(?=\*\*{next_label}\*\*:|\Z)"
        else:
            field_pattern = rf"\*\*{label}\*\*:\s*([\s\S]*?)\s*$"
        found = re.search(field_pattern, block)
        return found.group(1).strip() if found else ""

    return {
        "name": matched_name,
        "objective": bold_field("目的", "标准流程"),
        "procedure": bold_field("标准流程", "主要数据"),
        "data": bold_field("主要数据"),
        "matched": True,
    }


def build_standard_report_template(experiment_name: str) -> tuple[str, dict | None]:
    """Build a searchable standard report structure without inventing experiment facts."""

    requested = safe_text(experiment_name, "未命名实验")
    reference = search_experiment_reference(requested)
    resolved_name = reference["name"] if reference else requested
    objective = reference["objective"] if reference else "【待根据导师指导书填写】"
    procedure = reference["procedure"] if reference else "【待根据导师指导书填写】"
    data = reference["data"] if reference else "【待根据导师指导书定义数据项、单位和公式】"
    source_note = (
        "项目大学化学实验模块参考。只提供通用目的、流程和数据字段，不作为本校实验参数。"
        if reference
        else "未在项目实验目录中找到匹配项，仅生成标准报告结构。"
    )
    report = f"""# {safe_text(resolved_name)} - 标准实验报告模板

> 状态：NEEDS_HUMAN_REVIEW。{source_note}

## 实验目的

{_section(objective)}

## 实验原理

【待上传导师指导书后提取，不根据实验名称猜测公式或机理】

## 仪器与试剂

【待根据导师指导书填写仪器型号、试剂规格和实际用量】

## 安全与废物处理

【待根据指导书、SDS 和实验室规定填写】

## 计划实验步骤

{_section(procedure)}

## 实际操作与偏差

【实验完成后填写本次真实操作和计划偏差】

## 实验现象

【实验完成后填写真实观察，不得使用预期现象替代】

## 原始数据

主要数据字段：{_section(data)}

| 数据项 | 单位 | 原始值 | 备注 |
| --- | --- | --- | --- |
| 【按指导书拆分】 |  |  |  |

## 数据处理与作图

【待确认公式、单位、有效数字、拟合方法和专业软件要求】

## 实验结果

【待由本次真实数据和软件产物生成】

## 误差分析

【结合本次操作、仪器和数据证据填写，不使用通用套话】

## 讨论与思考题

【待根据导师问题和本次结果填写】

## 结论

【结论必须由本次真实数据支持】

## 人工审核

- [ ] 已上传并核对本次实验指导书
- [ ] 原始数据、单位和有效数字已核对
- [ ] 图表、参数表和正文来自同一批数据
- [ ] 缺失项已处理，最终内容已人工确认
"""
    return report, reference


def markdown_to_docx_bytes(markdown: str, status: str = "NEEDS_HUMAN_REVIEW") -> bytes:
    import sys

    scripts = ROOT / "scripts"
    if str(scripts) not in sys.path:
        sys.path.insert(0, str(scripts))
    from render_report_package import markdown_to_docx

    with tempfile.TemporaryDirectory() as temp_dir:
        path = Path(temp_dir) / "report.docx"
        markdown_to_docx(markdown, path, status)
        return path.read_bytes()
