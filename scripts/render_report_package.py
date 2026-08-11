"""Render pre-lab and formal chemistry reports from report-first JSON contracts."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

from validate_report_package import PROJECT_ROOT, validate_package


STATUS_LABELS = {
    "source_verified": "来源已核验",
    "user_provided": "用户提供",
    "software_imported": "软件导入",
    "model_draft": "模型草稿",
    "missing": "待填写",
    "needs_review": "待审核",
}


def load(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def display(value: Any, missing: str = "【待填写】") -> str:
    if value is None or value == "":
        return missing
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, float):
        return f"{value:.8g}"
    return str(value)


def status_text(value: str) -> str:
    return STATUS_LABELS.get(value, value)


def md_table(headers: list[str], rows: list[list[Any]]) -> list[str]:
    escaped = lambda value: display(value).replace("|", "\\|").replace("\n", "<br>")
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join("---" for _ in headers) + " |"]
    lines.extend("| " + " | ".join(escaped(value) for value in row) + " |" for row in rows)
    return lines


def render_prelab(spec: dict[str, Any], audit: dict[str, Any]) -> str:
    lines = [
        f"# {spec['experiment_name']} - 预习报告",
        "",
        f"> 报告状态：{audit['status']}。本报告只用于预习与记录设计；未确认参数不得直接用于真实实验操作。",
        "",
        "## 1. 实验基本信息与目的",
        "",
        f"实验 ID：`{spec['experiment_id']}`；协议版本：`{spec['schema_version']}`；配置状态：`{spec['status']}`。",
        "",
    ]
    for item in spec["objectives"]:
        lines.append(f"- {item['text']}（{status_text(item['content_status'])}）")

    lines.extend(["", "## 2. 实验原理", ""])
    for item in spec["principles"]:
        lines.extend([f"### {item['principle_id']}", "", item["text"], ""])
        if item["equations"]:
            lines.append("公式：" + "；".join(f"`{equation}`" for equation in item["equations"]))
        if item["assumptions"]:
            lines.append("适用假设：" + "；".join(item["assumptions"]))
        lines.append("来源：" + "、".join(item["evidence_refs"]))

    lines.extend(["", "## 3. 仪器、试剂与安全", ""])
    resource_rows: list[list[Any]] = []
    for group, resources in (("试剂", spec["materials"]), ("仪器", spec["instruments"])):
        for resource in resources:
            params = "; ".join(
                f"{p['name']}={display(p['value'])}{(' ' + p['unit']) if p['unit'] else ''} [{status_text(p['content_status'])}]"
                for p in resource["parameters"]
            ) or "记录实际规格/型号"
            resource_rows.append([group, resource["name"], resource["role"], params, status_text(resource["content_status"])])
    lines.extend(md_table(["类型", "名称", "用途", "参数/记录要求", "状态"], resource_rows))
    lines.extend(["", "### 安全与废物", ""])
    for label, values in (("个人防护", spec["safety"]["ppe"]), ("主要风险", spec["safety"]["hazards"]), ("废物处理", spec["safety"]["waste"]), ("应急", spec["safety"]["emergency"])):
        lines.append(f"- {label}：{'；'.join(values) if values else '【待指导书或教师确认】'}")

    lines.extend(["", "## 4. 实验设计与计划步骤", ""])
    step_rows = []
    for step in spec["planned_steps"]:
        params = "; ".join(
            f"{p['name']}={display(p['value'])}{(' ' + p['unit']) if p['unit'] else ''} [{status_text(p['content_status'])}]"
            for p in step["parameters"]
        ) or "无固定参数"
        step_rows.append([step["step_id"], step["stage"], step["action"], params, "; ".join(step["stop_conditions"]) or "无"])
    lines.extend(md_table(["步骤", "阶段", "计划操作", "关键参数", "停止条件"], step_rows))

    lines.extend(["", "## 5. 原始数据记录表", ""])
    lines.extend(md_table(
        ["字段 ID", "名称", "类型", "单位", "必填", "采集方式"],
        [[f["field_id"], f["label"], f["data_type"], f["unit"] or "-", f["required"], f["source_method"]] for f in spec["data_fields"]],
    ))

    lines.extend(["", "## 6. 数据处理与专业软件计划", ""])
    for task in spec["analysis_plan"]:
        lines.extend([
            f"### {task['name']}", "",
            f"处理等级：**{task['tier']}**；方法：`{task['method']}`；输入：{', '.join(task['input_field_ids']) or '待确认'}；要求输出：{', '.join(task['required_outputs']) or '待确认'}。",
        ])
        for step in task["guidance_steps"]:
            lines.append(f"- {step}")

    lines.extend(["", "## 7. 质量判断计划", ""])
    lines.extend(md_table(
        ["规则", "指标", "判定", "阈值", "失败动作", "状态"],
        [[r["rule_id"], r["metric"], r["operator"], r["threshold"], r["on_fail"], status_text(r["content_status"])] for r in spec["quality_rules"]],
    ))

    lines.extend(["", "## 8. 实验前人工确认", ""])
    for label, value in (
        ("指导书", spec["review"]["guide_confirmed"]),
        ("公式", spec["review"]["formulas_confirmed"]),
        ("质量规则", spec["review"]["quality_rules_confirmed"]),
        ("安全", spec["review"]["safety_confirmed"]),
    ):
        lines.append(f"- [{'x' if value else ' '}] {label}已确认")
    lines.extend(["", f"> 审核备注：{spec['review']['notes']}", "", "## 9. 参考来源", ""])
    for source in spec["sources"]:
        lines.append(f"- [{source['title']}]({source['location']})；类型：{source['source_type']}；版本：{display(source.get('version'), '未注明')}。")
    return "\n".join(lines) + "\n"


def summarize_values(values: dict[str, Any], field_map: dict[str, dict[str, Any]]) -> str:
    parts = []
    for key, value in values.items():
        field = field_map.get(key, {"label": key, "unit": None})
        unit = f" {field['unit']}" if field.get("unit") and value not in (None, "") else ""
        parts.append(f"{field['label']}={display(value)}{unit}")
    return "; ".join(parts)


def narrative_or_placeholder(items: list[dict[str, Any]], label: str) -> list[str]:
    available = [item for item in items if item.get("text")]
    if not available:
        return [f"> 【待填写：请根据真实实验记录补充{label}，Agent 不会编造。】"]
    return [f"- {item['text']}（{status_text(item['content_status'])}；证据：{', '.join(item['evidence_refs']) or '待补'}）" for item in available]


def render_formal(
    spec: dict[str, Any],
    run: dict[str, Any],
    artifacts: list[dict[str, Any]],
    audit: dict[str, Any],
) -> str:
    field_map = {item["field_id"]: item for item in spec["data_fields"]}
    step_map = {item["step_id"]: item for item in spec["planned_steps"]}
    lines = [
        f"# {spec['experiment_name']} - 实验报告",
        "",
        f"> 报告状态：{audit['status']}。阻断项 {audit['counts']['blocking']}，待审核项 {audit['counts']['warning']}。未完成最终人工审核前，本文件始终是草稿。",
        "",
        "## 报告元数据",
        "",
    ]
    participants = "、".join(display(p["name"], "匿名") + f"（{p['role']}）" for p in run["participants"]) or "【待填写】"
    metadata_rows = [
        ["实验/运行 ID", f"{spec['experiment_id']} / {run['run_id']}"],
        ["实验者", participants],
        ["时间", f"{display(run['started_at'])} 至 {display(run.get('ended_at'))}"],
        ["地点", display(run["location"]["value"])],
        ["协议/运行状态", f"ExperimentSpec {spec['schema_version']} / {run['run_status']}"],
        ["生成范围", "Agent 负责结构化整理、基础计算结果装配和一致性检查；真实数据、现象与最终结论由人工负责。"],
    ]
    lines.extend(md_table(["项目", "内容"], metadata_rows))

    lines.extend(["", "## 1. 实验预习与目的", ""])
    for item in spec["objectives"]:
        lines.append(f"- {item['text']}（来源状态：{status_text(item['content_status'])}）")
    for principle in spec["principles"]:
        lines.extend(["", principle["text"], "", "公式：" + "；".join(principle["equations"])])

    lines.extend(["", "## 2. 实验用品与实际条件", ""])
    if run["actual_conditions"]:
        lines.extend(md_table(
            ["条件", "实际值", "单位", "来源", "状态"],
            [[c["name"], c["value"], c["unit"] or "-", c["provenance"], status_text(c["content_status"])] for c in run["actual_conditions"]],
        ))
    else:
        lines.append("> 【待填写：仪器型号、试剂批次、环境及其他实际条件。】")

    lines.extend(["", "## 3. 实验设计", ""])
    lines.append("本次运行使用配置中定义的数据字段、分析计划和质量规则；任何与指导书不同的操作必须在实际步骤中记录偏差。")
    lines.append(f"原始记录数：{len(run['raw_records'])}；分析产物数：{len(artifacts)}。")

    lines.extend(["", "## 4. 实验前检查", ""])
    for label, value in (
        ("指导书", spec["review"]["guide_confirmed"]), ("公式", spec["review"]["formulas_confirmed"]),
        ("质量规则", spec["review"]["quality_rules_confirmed"]), ("安全", spec["review"]["safety_confirmed"]),
    ):
        lines.append(f"- [{'x' if value else ' '}] {label}已确认")

    lines.extend(["", "## 5. 实际实验操作", ""])
    actual_by_step = {item["planned_step_id"]: item for item in run["actual_steps"]}
    step_rows = []
    for step_id, planned in step_map.items():
        actual = actual_by_step.get(step_id)
        params = ""
        if actual:
            params = "; ".join(f"{p['name']}={display(p['value'])}{(' ' + p['unit']) if p['unit'] else ''}" for p in actual["actual_parameters"])
        step_rows.append([
            step_id,
            planned["action"],
            display(actual.get("actual_action") if actual else None),
            params or "【待记录】",
            display(actual.get("deviation") if actual else None, "无记录"),
        ])
    lines.extend(md_table(["步骤", "指导书计划", "实际执行", "实际参数", "偏差及原因"], step_rows))

    lines.extend(["", "## 6. 实验现象与原始数据", "", "### 实验现象", ""])
    observations = [item for item in run["observations"] if item.get("text")]
    if observations:
        lines.extend(md_table(
            ["编号", "步骤", "现象", "记录时间", "来源状态"],
            [[o["observation_id"], o["step_id"], o["text"], display(o["recorded_at"]), status_text(o["content_status"])] for o in observations],
        ))
    else:
        lines.append("> 【待填写：必须根据实验现场记录真实现象，禁止由 Agent 推测。】")
    lines.extend(["", "### 原始数据", ""])
    lines.extend(md_table(
        ["记录 ID", "数据", "来源文件", "记录时间", "状态"],
        [[r["record_id"], summarize_values(r["values"], field_map), display(r["source_file"], "手工录入/公开转录"), display(r["recorded_at"]), status_text(r["content_status"])] for r in run["raw_records"]],
    ))
    if run["corrections"]:
        lines.extend(["", "### 原始记录更正", ""])
        lines.extend(md_table(
            ["目标", "原值", "新值", "原因", "修改人/时间"],
            [[c["target_path"], c["old_value"], c["new_value"], c["reason"], f"{c['changed_by']} / {c['changed_at']}"] for c in run["corrections"]],
        ))

    lines.extend(["", "## 7. 数据预处理", ""])
    lines.append("原始数据保持不覆盖；空白、稀释、单位换算和异常处理只按分析产物登记的方法执行。")
    if not run["corrections"]:
        lines.append("本次运行未登记原始记录更正。")

    lines.extend(["", "## 8. 实验结果处理", ""])
    if not artifacts:
        lines.append("> 【待导入：尚无自动计算或专业软件分析产物。】")
    for artifact in artifacts:
        lines.extend([
            f"### {artifact['method_name']}", "",
            f"处理等级：**{artifact['tier']}**；软件：{artifact['software']['name']} {display(artifact['software']['version'], '')}；模型：{display(artifact['model'])}；状态：**{artifact['status']}**。",
            "",
        ])
        if artifact["parameters"]:
            lines.extend(md_table(
                ["参数", "值", "单位", "来源状态", "来源位置"],
                [[p["name"], p["value"], p["unit"] or "-", status_text(p["content_status"]), p["source_locator"]] for p in artifact["parameters"]],
            ))
        else:
            lines.append("> 【待导入参数表：Agent 不会替代专业软件生成复杂拟合结果。】")
        for plot in artifact["plots"]:
            plot_path = Path(plot["path"])
            resolved = plot_path if plot_path.is_absolute() else PROJECT_ROOT / plot_path
            if resolved.exists() and resolved.suffix.lower() in {".png", ".jpg", ".jpeg"}:
                lines.extend(["", f"![{artifact['method_name']}图表]({resolved.as_posix()})", ""])
            else:
                lines.append(f"图表文件：`{plot['path']}`（{plot['media_type']}）")

    guided = [task for task in spec["analysis_plan"] if task["tier"] in {"guided", "external"}]
    if guided:
        lines.extend(["", "### 专业软件操作指导", ""])
        for task in guided:
            lines.append(f"**{task['name']}**")
            for step in task["guidance_steps"]:
                lines.append(f"- {step}")

    lines.extend(["", "## 9. 结果质量判断", ""])
    lines.extend(md_table(
        ["规则", "指标", "阈值", "规则状态", "失败动作"],
        [[r["rule_id"], r["metric"], r["threshold"], status_text(r["content_status"]), r["on_fail"]] for r in spec["quality_rules"]],
    ))
    lines.extend(["", "### 一致性审查", ""])
    if audit["issues"]:
        lines.extend(md_table(
            ["级别", "代码", "问题", "位置"],
            [[i["severity"], i["code"], i["message"], i["path"]] for i in audit["issues"]],
        ))
    else:
        lines.append("未发现结构或追溯问题。")

    lines.extend(["", "## 10. 实验误差分析", ""])
    lines.extend(narrative_or_placeholder(run["error_analysis"], "有证据的误差分析"))
    lines.extend(["", "## 11. 讨论", ""])
    lines.extend(narrative_or_placeholder(run["discussion"], "结果讨论与方法限制"))
    lines.extend(["", "## 12. 实验总结", ""])
    conclusion = run["conclusion"]
    if conclusion.get("text"):
        lines.append(f"{conclusion['text']}（{status_text(conclusion['content_status'])}；证据：{', '.join(conclusion['evidence_refs']) or '待补'}）")
    else:
        lines.append("> 【待填写：结论必须依据真实结果、质量判断和实验目标，不得由 Agent 编造。】")

    lines.extend(["", "## 13. 清理与归档", ""])
    lines.append("实验废物和仪器清理应按本校指导书确认；本报告包归档 ExperimentSpec、ExperimentRun、分析产物、审查记录、Markdown 和 Word 文件。")

    lines.extend(["", "## 参考来源", ""])
    for source in spec["sources"]:
        lines.append(f"- [{source['title']}]({source['location']})；适用范围需按来源类型分别解释。")

    lines.extend(["", "## 人工审核记录", ""])
    final = run["final_review"]
    for label, key in (("原始数据", "raw_data_confirmed"), ("实际步骤", "actual_steps_confirmed"), ("分析产物", "analysis_confirmed"), ("最终报告", "report_confirmed")):
        lines.append(f"- [{'x' if final[key] else ' '}] {label}已确认")
    lines.extend(["", f"审核人：{display(final['reviewer'])}；审核时间：{display(final['reviewed_at'])}。", "", f"> 审核备注：{final['notes']}"])
    return "\n".join(lines) + "\n"


def clean_markdown(text: str) -> str:
    text = re.sub(r"!\[[^]]*\]\([^)]+\)", "", text)
    text = re.sub(r"\[([^]]+)\]\([^)]+\)", r"\1", text)
    return text.replace("**", "").replace("`", "")


def parse_md_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        values = [clean_markdown(cell.strip().replace("<br>", "\n")) for cell in lines[index].strip().strip("|").split("|")]
        if index != start + 1:
            rows.append(values)
        index += 1
    return rows, index


def markdown_to_docx(markdown: str, output: Path, status: str) -> None:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    def set_font(run: Any, size: float, bold: bool = False, color: str = "000000", italic: bool = False) -> None:
        run.font.name = "Calibri"
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        run.font.color.rgb = RGBColor.from_string(color)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.add_run("化学实验报告辅助 Agent"), 8.5, color="667085")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("报告状态：" + status + "  |  第 "), 8, color="667085")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)
    set_font(footer.add_run(" 页"), 8, color="667085")

    numbering = doc.part.numbering_part.element
    next_abs = max([int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))] or [0]) + 1
    next_num = max([int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))] or [0]) + 1

    def make_numbering(fmt: str, text: str, abs_id: int, num_id: int) -> int:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abs_id))
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); level.append(start)
        num_fmt = OxmlElement("w:numFmt"); num_fmt.set(qn("w:val"), fmt); level.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText"); lvl_text.set(qn("w:val"), text); level.append(lvl_text)
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab"); tab.set(qn("w:val"), "num"); tab.set(qn("w:pos"), "540"); tabs.append(tab); ppr.append(tabs)
        ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "540"); ind.set(qn("w:hanging"), "270"); ppr.append(ind)
        spacing = OxmlElement("w:spacing"); spacing.set(qn("w:after"), "80"); spacing.set(qn("w:line"), "300"); spacing.set(qn("w:lineRule"), "auto"); ppr.append(spacing)
        level.append(ppr); abstract.append(level); numbering.append(abstract)
        num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id)); abstract_ref = OxmlElement("w:abstractNumId"); abstract_ref.set(qn("w:val"), str(abs_id)); num.append(abstract_ref); numbering.append(num)
        return num_id

    bullet_num = make_numbering("bullet", "•", next_abs, next_num)
    decimal_num = make_numbering("decimal", "%1.", next_abs + 1, next_num + 1)

    def add_list(text: str, num_id: int) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0"); num_pr.append(ilvl)
        nid = OxmlElement("w:numId"); nid.set(qn("w:val"), str(num_id)); num_pr.append(nid)
        set_font(p.add_run(clean_markdown(text)), 10.5)

    def set_cell(cell: Any, width: int, header_cell: bool = False) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.find(qn("w:tcW"))
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
        tc_w.set(qn("w:w"), str(width)); tc_w.set(qn("w:type"), "dxa")
        if tc_w.getparent() is None: tc_pr.append(tc_w)
        margins = tc_pr.find(qn("w:tcMar"))
        if margins is None:
            margins = OxmlElement("w:tcMar")
        if margins.getparent() is None: tc_pr.append(margins)
        for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
            node = margins.find(qn(f"w:{side}"))
            if node is None:
                node = OxmlElement(f"w:{side}")
            node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")
            if node.getparent() is None: margins.append(node)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if header_cell:
            shade = OxmlElement("w:shd"); shade.set(qn("w:fill"), "E8EEF5"); tc_pr.append(shade)

    def add_table(rows: list[list[str]]) -> None:
        if not rows:
            return
        cols = max(len(row) for row in rows)
        max_chars = [max(4, max((len(row[i]) if i < len(row) else 0) for row in rows)) for i in range(cols)]
        raw = [max(900, min(3600, size * 135)) for size in max_chars]
        scale = 9360 / sum(raw)
        widths = [int(value * scale) for value in raw]
        widths[-1] += 9360 - sum(widths)
        table = doc.add_table(rows=len(rows), cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        tbl_pr = table._tbl.tblPr
        layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed"); tbl_pr.append(layout)
        tbl_w = tbl_pr.find(qn("w:tblW")); tbl_w.set(qn("w:w"), "9360"); tbl_w.set(qn("w:type"), "dxa")
        tbl_ind = OxmlElement("w:tblInd"); tbl_ind.set(qn("w:w"), "120"); tbl_ind.set(qn("w:type"), "dxa"); tbl_pr.append(tbl_ind)
        grid = table._tbl.tblGrid
        for child in list(grid): grid.remove(child)
        for width in widths:
            col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
        for row_index, values in enumerate(rows):
            for col_index in range(cols):
                cell = table.cell(row_index, col_index)
                cell.text = values[col_index] if col_index < len(values) else ""
                set_cell(cell, widths[col_index], row_index == 0)
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.15
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_index == 0 or len(cell.text) < 12 else WD_ALIGN_PARAGRAPH.LEFT
                    for run in p.runs:
                        set_font(run, 8.5 if cols >= 5 else 9, bold=row_index == 0)
        tr_pr = table.rows[0]._tr.get_or_add_trPr(); repeat = OxmlElement("w:tblHeader"); repeat.set(qn("w:val"), "true"); tr_pr.append(repeat)
        after = doc.add_paragraph(); after.paragraph_format.space_after = Pt(2)

    lines = markdown.splitlines()
    index = 0
    first_title = True
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and re.match(r"^\|[\s:|-]+\|$", lines[index + 1].strip()):
            rows, index = parse_md_table(lines, index)
            add_table(rows)
            continue
        image_match = re.match(r"!\[([^]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            image_path = Path(image_match.group(2))
            if image_path.exists():
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                inline_shape = p.add_run().add_picture(str(image_path), width=Inches(5.7))
                inline_shape._inline.docPr.set("descr", image_match.group(1) or image_path.name)
                inline_shape._inline.docPr.set("title", image_match.group(1) or "实验图表")
                caption = doc.add_paragraph(); caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_font(caption.add_run("图：" + image_match.group(1)), 9, color="667085")
            index += 1
            continue
        if stripped.startswith("# ") and first_title:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
            set_font(p.add_run(clean_markdown(stripped[2:])), 22, bold=True, color="0B2545")
            sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER; sub.paragraph_format.space_after = Pt(14)
            set_font(sub.add_run("标准化化学实验报告 | 可追溯草稿"), 10, color="667085")
            first_title = False
        elif stripped.startswith("### "):
            doc.add_paragraph(clean_markdown(stripped[4:]), style="Heading 2")
        elif stripped.startswith("## "):
            doc.add_paragraph(clean_markdown(stripped[3:]), style="Heading 1")
        elif stripped.startswith("# "):
            doc.add_paragraph(clean_markdown(stripped[2:]), style="Heading 1")
        elif stripped.startswith("> "):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.18); p.paragraph_format.right_indent = Inches(0.12)
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(8)
            p_pr = p._p.get_or_add_pPr(); borders = OxmlElement("w:pBdr"); left = OxmlElement("w:left"); left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18"); left.set(qn("w:color"), "B7791F"); left.set(qn("w:space"), "8"); borders.append(left); p_pr.append(borders)
            set_font(p.add_run(clean_markdown(stripped[2:])), 9.5, color="7A5A00")
        elif re.match(r"^- \[[ xX]\] ", stripped):
            mark = "已确认" if stripped[3].lower() == "x" else "未确认"
            add_list(f"{mark}：{stripped[6:]}", bullet_num)
        elif stripped.startswith("- "):
            add_list(stripped[2:], bullet_num)
        elif re.match(r"^\d+\. ", stripped):
            add_list(re.sub(r"^\d+\. ", "", stripped), decimal_num)
        else:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_font(p.add_run(clean_markdown(stripped)), 10.5)
        index += 1

    doc.core_properties.title = clean_markdown(lines[0][2:]) if lines and lines[0].startswith("# ") else "化学实验报告"
    doc.core_properties.subject = "标准化化学实验报告辅助 Agent 输出"
    doc.core_properties.author = "化学实验报告辅助 Agent"
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--spec", type=Path, required=True)
    parser.add_argument("--run", type=Path, required=True)
    parser.add_argument("--artifact", type=Path, action="append", default=[])
    parser.add_argument("--output-dir", type=Path, required=True)
    args = parser.parse_args()
    spec = load(args.spec)
    run = load(args.run)
    artifacts = [load(path) for path in args.artifact]
    audit = validate_package(spec, run, artifacts, PROJECT_ROOT)
    if audit["status"] == "BLOCKED":
        raise ValueError("报告包存在阻断问题，请先运行 validate_report_package.py 查看。")
    args.output_dir.mkdir(parents=True, exist_ok=True)
    prelab_md = render_prelab(spec, audit)
    formal_md = render_formal(spec, run, artifacts, audit)
    paths = {
        "prelab_md": args.output_dir / "预习报告.md",
        "formal_md": args.output_dir / "实验报告_草稿.md",
        "prelab_docx": args.output_dir / "预习报告.docx",
        "formal_docx": args.output_dir / "实验报告_草稿.docx",
        "audit": args.output_dir / "报告审查.json",
    }
    paths["prelab_md"].write_text(prelab_md, encoding="utf-8")
    paths["formal_md"].write_text(formal_md, encoding="utf-8")
    paths["audit"].write_text(json.dumps(audit, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    markdown_to_docx(prelab_md, paths["prelab_docx"], audit["status"])
    markdown_to_docx(formal_md, paths["formal_docx"], audit["status"])
    for label, path in paths.items():
        print(f"{label.upper()} {path}")
    print(f"STATUS {audit['status']}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
